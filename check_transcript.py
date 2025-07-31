from docx import Document
import os

def detailed_transcript_analysis():
    file_path = 'trans.docx'
    
    print(f"=== ДЕТАЛЬНЫЙ АНАЛИЗ {file_path} ===")
    
    if not os.path.exists(file_path):
        print(f"❌ Файл не найден!")
        return
    
    try:
        doc = Document(file_path)
        
        print(f"📄 Общее количество параграфов в документе: {len(doc.paragraphs)}")
        
        all_text = ""
        non_empty_paragraphs = []
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            all_text += para.text + " "
            
            if text:
                non_empty_paragraphs.append((i, text))
        
        print(f"📝 Параграфов с текстом: {len(non_empty_paragraphs)}")
        print(f"📊 Общая длина всего текста: {len(all_text)} символов")
        print(f"📊 Длина текста без пробелов: {len(all_text.replace(' ', ''))}")
        
        if len(non_empty_paragraphs) == 0:
            print("❌ ПРОБЛЕМА: В документе НЕТ текста!")
            return
        
        print(f"\n📄 ВСЕ параграфы с текстом:")
        for i, (para_num, text) in enumerate(non_empty_paragraphs):
            print(f"{i+1}. [Параграф {para_num}]: {text[:150]}{'...' if len(text) > 150 else ''}")
        
        # Проверяем, достаточно ли текста для анализа
        total_text = ' '.join([text for _, text in non_empty_paragraphs])
        sentences = [s.strip() for s in total_text.replace('!', '.').replace('?', '.').split('.') if s.strip()]
        
        print(f"\n📊 Количество предложений: {len(sentences)}")
        print(f"📊 Средняя длина предложения: {len(total_text) / len(sentences) if sentences else 0:.1f} символов")
        
        if len(sentences) < 5:
            print("⚠️  ВНИМАНИЕ: Очень мало текста для качественного анализа!")
        
        if len(total_text) < 100:
            print("⚠️  ВНИМАНИЕ: Текст слишком короткий для анализа компетенций!")
            
    except Exception as e:
        print(f"❌ Ошибка при чтении: {e}")

if __name__ == "__main__":
    detailed_transcript_analysis() 